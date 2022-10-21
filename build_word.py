# 导入库
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn  # 中文字体
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER  # 会有红色下划线报异常，不过可以正常使用
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from ExcelOp import ExcelOp

moduleType = ['基础数据', 'AI', 'SFA', 'DMS', 'TPM', 'PMM', '其他']

# 新建空白文档
from PG import PgInfo

doc1 = Document()
doc1.styles['Normal'].font.name = '宋体'  # 设置字体
doc1.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

excel_op = ExcelOp('/Users/lin/PycharmProjects/pythonProject/xuanwu/data/智慧100 数据字典整理.xlsx')
list1 = excel_op.get_col_value(5)
tableType = set(list1)
print(tableType)
dic = {}

doc1.add_heading('数据字典整理', 0)
doc1.add_heading('SFA模块', 1)
pg = PgInfo('tenant_1000061')


tablename = 'kx_cost_customerapply'
sql = '''SELECT 
	A.attname AS field,
	T.typname AS 字段类型,
	A.attlen AS 字段最大容量,
	concat_ws (
		'',
		t.typname,
		SUBSTRING (
			format_type (a.atttypid, a.atttypmod)
			from
				'\(.*\)'
		)
	) as type,
	CASE A.attnotnull 
		WHEN 'f' THEN '否'
		WHEN 't' THEN '是' END AS 可否为空,
	b.description AS 注释 
FROM
	pg_class C,
	pg_attribute A LEFT OUTER JOIN pg_description b ON A.attrelid = b.objoid 
	AND A.attnum = b.objsubid,
	pg_type T 
WHERE
	C.relname = ''' + "'{0}'".format(tablename) + ''' AND A.attnum > 0 
AND A.attrelid = C.oid 
AND A.atttypid = T.oid 
ORDER BY
A.attnum'''
list2 = pg.ExecQuery_And_Get_Return(sql)
row = len(list2)
col = len(list2[0])
table = doc1.add_table(rows=row, cols=col, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

table.autofit = True
for i in range(len(list2)):
    rowCells = table.rows[i].cells
    for j in range(6):
        rowCells[j].text = str(list2[i][j])

doc1.save('w2.docx')
