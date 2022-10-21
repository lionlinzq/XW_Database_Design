import json

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from ExcelOp import ExcelOp
from PG import PgInfo
from docx.oxml.ns import qn

excelPath = "/Users/lin/PycharmProjects/pythonProject/xuanwu/data/智慧100 数据字典整理.xlsx"

excel_op = ExcelOp(excelPath)
# 一级模块
firstType = excel_op.get_col_value(5)
# 二级模块
secondType = excel_op.get_col_value(6)
# 对应的表
tableType = excel_op.get_col_value(1)
# 对应表名
tableName = excel_op.get_col_value(2)
table_name = {}
for i in range(len(tableType)):
    table_name[tableType[i]] = tableName[i]

moduleType = ['AI', 'SFA', 'DMS', 'TPM', 'PMM', '基础数据', '报表', '其他']
fieldType = ['序号', '字段名称', '类型(长度)', '可否为空', '注释']

# 按照一级模块存储
moduleFirst = {}
# 按照二级模块存储
moduleSecond = {}
# 二级模块存储表名
moduleSecondList = []

pg = PgInfo('tenant_1000061')
doc1 = Document()
doc1.add_heading('数据字典整理', 0)
doc1.styles['Normal'].font.name = '宋体'  # 设置字体
doc1.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
doc1.styles['Normal'].font.size = Pt(10.5)


# 获取还在数据库中的表
def get_noUse_tables():
    sql = '''SELECT relname FROM pg_class C,pg_namespace n WHERE 
    C.relnamespace = n.oid AND nspname = 'public' AND relkind = 'r' ORDER BY relname'''
    tableList = pg.ExecQuery_And_Get_Return(sql)
    usedTableSet = set()
    for tables in tableList:
        usedTableSet.add(tables[0])
    excel_tables = set(tableType)
    # 获取到excel中有但是数据库已经没有的表
    noUseTables = excel_tables - usedTableSet
    return noUseTables


def get_excel_use_tables():
    listA = excel_op.get_all_rows_value()
    print(listA)


# 根据文档生成数据表结构信息,字典形式存储
def make_table_structure():
    noUseTables = get_noUse_tables()
    for i in range(len(firstType)):
        moduleSecondList = []
        if i == 0:
            continue
        first, second, table = firstType[i], secondType[i], tableType[i]
        if table in noUseTables:
            continue
        if first is not None:
            if second is not None:
                if second in moduleSecond.keys():
                    moduleSecondList = moduleSecond[second]
                    moduleSecondList.append(table)
                    moduleSecond[second] = moduleSecondList
                    moduleFirst[first][0][second] = moduleSecond[second]
                else:
                    moduleSecondList.append(table)
                    moduleSecond[second] = moduleSecondList
                    if first in moduleFirst.keys():
                        moduleFirst[first][0][second] = moduleSecond[second]
                        continue
                    moduleFirst[first] = [{second: moduleSecond[second]}]
            else:
                second = first
                if second in moduleSecond.keys():
                    moduleSecondList = moduleSecond[second]
                    moduleSecondList.append(table)
                    moduleSecond[second] = moduleSecondList
                    moduleFirst[first][0][second] = moduleSecond[second]
                else:
                    moduleSecondList.append(table)
                    moduleSecond[second] = moduleSecondList
                    if first in moduleFirst.keys():
                        moduleFirst[first][0][second] = moduleSecond[second]
                        continue
                    moduleFirst[first] = [{second: moduleSecond[second]}]

    print(moduleFirst)
    j = json.dumps(moduleFirst, ensure_ascii=False)
    print(j)
    return moduleFirst


# 生成每张表的信息并保存进一个表格中
def make_table_info_and_build_table(tablename):
    sql = '''SELECT 
    	A.attname AS field,
    	A.attlen AS 字段最大容量,
    	concat_ws (
    		'',
    		t.typname,
    		SUBSTRING (
    			format_type (a.atttypid, a.atttypmod)
    			from
    				'\(.*\)'
    		)
    	) as type,
    	CASE A.attnotnull 
    		WHEN 'f' THEN '否'
    		WHEN 't' THEN '是' END AS 可否为空,
    	b.description AS 注释 
    FROM
    	pg_class C,
    	pg_attribute A LEFT OUTER JOIN pg_description b ON A.attrelid = b.objoid 
    	AND A.attnum = b.objsubid,
    	pg_type T 
    WHERE
    	C.relname = ''' + "'{0}'".format(tablename) + ''' AND A.attnum > 0 
    AND A.attrelid = C.oid 
    AND A.atttypid = T.oid 
    ORDER BY
    A.attnum'''
    list2 = pg.ExecQuery_And_Get_Return(sql)
    row = len(list2)
    col = len(list2[0])
    table = doc1.add_table(rows=row, cols=col, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i in range(len(list2)):
        rowCells = table.rows[i].cells
        for j in range(6):
            rowCells[j].text = str(list2[i][j])


# 设置表格第一行为标题
def build_word():
    # moduleType = ['SFA', 'DMS', 'TPM', 'PMM', '平台表', '多模块使用', '报表']
    table_dictionary = make_table_structure()
    # for firstModule in table_dictionary:
    for firstModule in moduleType:
        # firstModule: SFA、DMS...
        # secondModules: {'DMS': ['kx_sale_target'], 'DMS-订单设计': ['kx_return_order_detail']}
        secondModules = table_dictionary[firstModule][0]
        doc1.add_heading('{0}模块'.format(firstModule), 1)
        for secondModule in secondModules:
            # secondModule: SFA-拜访计划制定、SFA-考勤管理
            doc1.add_heading(secondModule, 2)
            tableList = secondModules[secondModule]
            b = 1
            for table in tableList:
                doc1.add_heading('{0}.{1} {2}'.format(b, table, table_name[table]), 3)
                b = b + 1
                # 查询表主键sql
                keySql = '''select pg_attribute.attname as colname
                            from
                            pg_constraint  inner join pg_class
                            on pg_constraint.conrelid = pg_class.oid
                            inner join pg_attribute on pg_attribute.attrelid = pg_class.oid
                            and  pg_attribute.attnum = pg_constraint.conkey[1]
                            where pg_class.relname = '{0}'
                            and pg_constraint.contype='p' '''.format(table)
                keyList = pg.ExecQuery_And_Get_Return(keySql)
                # print(b, keyList)
                if keyList:
                    doc1.add_paragraph('主键:{}'.format(keyList[0][0]))
                # 查询索引sql
                selectIndexSql = '''select indexname from pg_indexes where tablename = ''' + "'{}'".format(table)
                if table == 'sys_html_function':
                    print(selectIndexSql)
                indexDict = pg.ExecQuery_And_Get_Return(selectIndexSql)
                if indexDict:
                    indexList = []
                    for index in range(len(indexDict)):
                        indexList.append("{0}.{1}".format(index + 1, indexDict[index][0]))
                    # print(table)
                    indexStr = u"索引:{0}".format('、'.join(indexList))
                    print(indexStr)
                    doc1.add_paragraph(indexStr)
                # 查询字段信息sql
                sql = '''SELECT
                                A.attname AS field,
                                concat_ws (
                                    '',
                                    t.typname,
                                    SUBSTRING (
                                        format_type (a.atttypid, a.atttypmod)
                                        from
                                            '\(.*\)'
                                    )
                                ) as type,
                                CASE A.attnotnull
                                    WHEN 'f' THEN '否'
                                    WHEN 't' THEN '是' END AS 可否为空,
                                b.description AS 注释
                            FROM
                                pg_class C,
                                pg_attribute A LEFT OUTER JOIN pg_description b ON A.attrelid = b.objoid
                                AND A.attnum = b.objsubid,
                                pg_type T
                            WHERE
                                C.relname = ''' + "'{0}'".format(
                    table) + ''' AND A.attnum > 0  AND A.attrelid = C.oid AND A.atttypid = T.oid  ORDER BY A.attnum'''
                list2 = pg.ExecQuery_And_Get_Return(sql)
                row = len(list2)
                col = len(list2[0])
                table = doc1.add_table(rows=row + 1, cols=col + 1, style='Table Grid')
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True
                for x in range(5):
                    run = table.cell(0, x).paragraphs[0].add_run(fieldType[x])
                    run.bold = True  # 加粗
                for i in range(len(list2)):
                    rowCells = table.rows[i + 1].cells
                    for j in range(len(list2[0])):
                        if j == 0:
                            rowCells[j].text = str(i + 1)
                        rowCells[j + 1].text = str(list2[i][j])
    doc1.save('数据字典整理示例.docx')


if __name__ == '__main__':
    build_word()
    # get_excel_use_tables()
